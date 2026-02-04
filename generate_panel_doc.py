#!/usr/bin/env python3
"""
Generate Evidify Platform Architecture & Research Thesis document.
Comprehensive technical and strategic analysis for expert panel review.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import shutil
import os

def set_cell_shading(cell, color_hex):
    """Set cell background shading."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table_row(table, cells_data, bold=False, header=False):
    """Add a row to a table with consistent formatting."""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        if bold or header:
            run.bold = True
        if header:
            set_cell_shading(cell, "2B3544")
            run.font.color.rgb = RGBColor(255, 255, 255)

def create_document():
    doc = Document()

    # ── STYLES ──────────────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 5):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Calibri'
        heading_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        if level == 1:
            heading_style.font.size = Pt(18)
            heading_style.paragraph_format.space_before = Pt(24)
            heading_style.paragraph_format.space_after = Pt(12)
        elif level == 2:
            heading_style.font.size = Pt(14)
            heading_style.paragraph_format.space_before = Pt(18)
            heading_style.paragraph_format.space_after = Pt(8)
        elif level == 3:
            heading_style.font.size = Pt(12)
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)

    # ── FOOTER ──────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("© 2026 Evidify. All rights reserved. Confidential — for expert review only.")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = 'Calibri'

    # ── TITLE PAGE ──────────────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Evidify Platform Architecture\n& Research Thesis")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    run.bold = True
    run.font.name = 'Calibri'

    doc.add_paragraph("")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Comprehensive Technical & Strategic Analysis\nfor Multidisciplinary Expert Panel Review")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = 'Calibri'

    doc.add_paragraph("")
    doc.add_paragraph("")

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run("Joshua M. Henderson, PhD")
    run.font.size = Pt(14)
    run.bold = True
    run.font.name = 'Calibri'

    role = doc.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = role.add_run("Research Platform Architect")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = 'Calibri'

    doc.add_paragraph("")

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date.add_run("February 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = 'Calibri'

    doc.add_paragraph("")

    classification = doc.add_paragraph()
    classification.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = classification.add_run("CONFIDENTIAL — FOR EXPERT REVIEW ONLY")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
    run.bold = True
    run.font.name = 'Calibri'

    doc.add_page_break()

    # ── TABLE OF CONTENTS ───────────────────────────────────────────────────
    doc.add_heading("Table of Contents", level=1)

    toc_items = [
        ("1.", "Executive Summary", 3),
        ("2.", "The Legal Problem", 4),
        ("3.", "Technical Architecture", 6),
        ("", "3a. Event System", 6),
        ("", "3b. Hash Chain (CanonicalHash.ts)", 7),
        ("", "3c. Impression Ledger", 8),
        ("", "3d. Export Pack", 9),
        ("", "3e. Derived Metrics", 10),
        ("", "3f. Local-First Architecture", 11),
        ("4.", "The Contrast — STI Prevention Thesis", 12),
        ("5.", "Clinical Workflow Integration", 14),
        ("6.", "Workload Monitoring", 16),
        ("7.", "What's Not Built Yet — Honest Gaps", 17),
        ("8.", "Research Design — Proposed BRPLL Collaboration", 19),
        ("9.", "Market & Strategic Position", 21),
        ("10.", "Vision — What This Could Become", 22),
        ("11.", "Questions for the Panel", 24),
    ]

    for num, title_text, page in toc_items:
        p = doc.add_paragraph()
        if num and num[0].isdigit() and "." == num[-1]:
            run = p.add_run(f"{num} {title_text}")
            run.bold = True
            run.font.size = Pt(11)
        else:
            run = p.add_run(f"    {num} {title_text}".rstrip())
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.font.name = 'Calibri'

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 1: EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("1. Executive Summary", level=1)

    doc.add_paragraph(
        "Evidify is a local-first research instrumentation platform that creates "
        "cryptographically chained, temporally ordered documentation of clinical "
        "decision-making in AI-assisted radiology. It is designed to produce "
        "court-defensible records proving that a human clinician exercised "
        "independent judgment before, during, and after AI consultation — and to "
        "measure the degree to which AI influenced that judgment."
    )

    doc.add_heading("The Core Problem", level=3)
    doc.add_paragraph(
        "AI-assisted clinical decision-making creates a liability vacuum. When a "
        "radiologist reads a mammogram with AI assistance and the outcome is adverse, "
        "the fundamental question in litigation is: did the radiologist exercise "
        "independent clinical judgment, or did they rubber-stamp the AI's output? "
        "Current clinical documentation systems — PACS timestamps, dictation logs, "
        "EHR entries — record what was decided but not how it was decided. They cannot "
        "prove temporal ordering of human judgment relative to AI consultation. A bare "
        "timestamp of '2 minutes, 14 seconds' invites jurors to make snap character "
        "judgments about the radiologist's diligence without understanding what those "
        "minutes contained."
    )

    doc.add_heading("The Thesis", level=3)
    doc.add_paragraph(
        "Evidify's thesis is that a cryptographically chained, append-only event log — "
        "where each entry incorporates a SHA-256 hash of the previous entry — can satisfy "
        "the requirements of self-authenticating digital evidence under Federal Rules of "
        "Evidence 902(13) and 902(14), while simultaneously providing the behavioral data "
        "needed to measure cognitive engagement through metrics such as pre-AI reading time, "
        "post-AI deliberation time, decision change patterns, and deviation documentation "
        "quality. The hash chain proves temporal ordering; the derived metrics prove "
        "cognitive engagement; the narrative documentation prevents Spontaneous Trait "
        "Inference by jurors."
    )

    doc.add_heading("Current State", level=3)
    doc.add_paragraph(
        "What exists today is a research preview — a functional demonstration of the "
        "instrumentation architecture, not a production clinical tool. The frontend "
        "(React/TypeScript, v4.3.0-beta) implements the complete event logging pipeline, "
        "hash chain construction, export packaging, and visualization components. The "
        "backend (Rust/Tauri) provides encrypted local storage and audit logging but is "
        "not yet integrated with the research instrumentation layer. No real mammogram "
        "images are loaded (placeholder dimensions are used). No real EHR or PACS "
        "integration exists. No IRB approval has been obtained. No attorney has formally "
        "reviewed the FRE 902 compliance claims. The platform has not been deployed in "
        "any clinical setting. This document describes what is built, what is planned, "
        "and where the gaps are."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 2: THE LEGAL PROBLEM
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("2. The Legal Problem", level=1)

    doc.add_heading("2.1 The Malpractice Liability Gap in AI-Assisted Radiology", level=2)
    doc.add_paragraph(
        "The integration of AI into diagnostic radiology creates a novel liability "
        "landscape. As referenced in the codebase (contrastDemoData.ts, TheContrast.tsx, "
        "and the one-pager document), Bernstein et al. (2025, Nature Health) have documented "
        "the emerging uncertainty around who is responsible when AI-assisted diagnoses go "
        "wrong. The fundamental tension: radiologists are expected to use AI tools that may "
        "improve diagnostic accuracy, but the act of consulting AI creates a record that can "
        "be weaponized in litigation regardless of outcome."
    )

    doc.add_paragraph(
        "If the radiologist agrees with AI and the diagnosis is wrong, the plaintiff argues "
        "the radiologist failed to exercise independent judgment. If the radiologist "
        "disagrees with AI and the diagnosis is wrong, the plaintiff argues the radiologist "
        "ignored a tool that would have caught the error. The radiologist cannot win unless "
        "the documentation proves the reasoning process, not just the outcome."
    )

    doc.add_heading("2.2 How Current Documentation Fails: Spontaneous Trait Inference", level=2)
    doc.add_paragraph(
        "TheContrast.tsx (frontend/src/components/research/legal/TheContrast.tsx) implements "
        "a side-by-side demonstration of this failure mode. The component draws directly from "
        "two bodies of research:"
    )

    p = doc.add_paragraph()
    run = p.add_run("Pennington & Hastie (1992), Story Model for Juror Decision Making: ")
    run.bold = True
    run.font.name = 'Calibri'
    run = p.add_run(
        "Jurors construct narrative accounts of events and evaluate evidence by how well "
        "it fits a coherent story. Bare timestamps without context fail to provide a story, "
        "leaving jurors to construct their own — typically unfavorable to the defendant."
    )

    p = doc.add_paragraph()
    run = p.add_run("Uleman, Saribay & Gonzalez (2008), Spontaneous Trait Inference (STI): ")
    run.bold = True
    run.font.name = 'Calibri'
    run = p.add_run(
        "Observers automatically and involuntarily infer character traits from behavioral "
        "descriptions. When a jury sees '2 minutes, 14 seconds on the case,' they "
        "involuntarily conclude 'careless' or 'rushed.' This inference occurs without "
        "conscious deliberation and is resistant to correction once formed."
    )

    doc.add_paragraph(
        "TheContrast.tsx implements this as a two-panel demonstration: the left panel "
        "('Prosecution Exhibit') shows a structured activity log with raw timestamps and "
        "durations. The right panel ('Defense Exhibit') presents the same data embedded in "
        "a clinical reasoning narrative. The component includes a 'highlight mode' that "
        "marks identical data points across both panels, demonstrating that the factual "
        "content is unchanged — only the framing differs. The research context section "
        "explains the STI mechanism and cites the relevant literature."
    )

    doc.add_heading("2.3 The 'Rubber Stamp' Problem", level=2)
    doc.add_paragraph(
        "ExpertWitnessExport.tsx (frontend/src/components/research/legal/ExpertWitnessExport.tsx) "
        "implements six specific rubber-stamp detection indicators, each with severity "
        "classification (low/medium/high):"
    )

    # Rubber stamp indicators table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ["Indicator", "Detection Logic", "Severity", "Threshold"], header=True)
    table.rows[0].cells[0].text = ""  # Clear auto-generated row

    indicators = [
        ("MINIMAL_PRE_AI_TIME", "Pre-AI reading time < threshold", "High if <5s, Medium if <15s", "≥15s recommended"),
        ("MINIMAL_POST_AI_TIME", "Post-AI review time < threshold", "High if <3s, Medium if <5s", "≥5s recommended"),
        ("INSTANT_AI_AGREEMENT", "Changed assessment to match AI with minimal review time (<5s post-AI)", "High", "Not triggered if no change"),
        ("UNDOCUMENTED_DEVIATION", "Disagreed with AI without documenting clinical rationale", "High", "Deviation doc required"),
        ("UNACKNOWLEDGED_FINDINGS", "AI-flagged regions not explicitly reviewed", "High", "All findings must be acknowledged"),
        ("DISCLOSURE_NOT_VIEWED", "AI performance disclosure was not viewed", "Medium", "Disclosure must be shown"),
    ]

    for ind in indicators:
        row = table.add_row()
        for i, val in enumerate(ind):
            row.cells[i].text = val
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.name = 'Calibri'

    doc.add_paragraph("")
    doc.add_paragraph(
        "The calculateRubberStampRiskLevel function classifies overall risk: ≥2 high-severity "
        "detections = 'high' risk; ≥1 high or ≥2 medium = 'medium'; otherwise 'low'. This "
        "classification feeds directly into the expert witness packet summary."
    )

    doc.add_heading("2.4 FRE 902(13)/(14) and Self-Authenticating Digital Evidence", level=2)
    doc.add_paragraph(
        "Federal Rules of Evidence 902(13) and 902(14) allow certified records generated by "
        "an electronic process to be self-authenticating — meaning they can be admitted without "
        "extrinsic evidence of authenticity if accompanied by a certification from a qualified "
        "person. The Evidify hash chain is designed with this standard in mind:"
    )

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Each event is hashed using SHA-256 over its canonical JSON payload (RFC 8785 compliant serialization)")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Each chain hash incorporates a fixed 128-byte structured input: sequence number (4B) + previous hash (32B) + event ID (36B) + timestamp (24B) + content hash (32B)")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("The genesis block uses 64 hex zeros as the previous hash")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Any modification to any entry invalidates all subsequent chain hashes, making tampering detectable and precisely locatable")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("The manifest includes a timestampTrustModel field explicitly set to 'client_clock_untrusted' — acknowledging that client-side timestamps are instrumentation data, not attestation")

    doc.add_paragraph(
        "The codebase also references Vermont's 12 V.S.A. §1913, a state-level statute governing "
        "digital evidence admissibility. The ExpertWitnessExport component's temporal proof "
        "walkthrough section includes plain-language explanation of how the hash chain meets "
        "self-authentication requirements, with a tamper detection demonstration that visually "
        "shows what happens when a chain entry is altered."
    )

    p = doc.add_paragraph()
    run = p.add_run("FOR THE LEGAL EXPERTS: ")
    run.bold = True
    run.font.name = 'Calibri'
    run = p.add_run(
        "The FRE 902 compliance claim is an architectural aspiration, not a legal opinion. "
        "No attorney has formally reviewed whether this specific implementation meets the "
        "certification requirements of 902(13)/(14). The 'client_clock_untrusted' trust model "
        "is an honest acknowledgment that without a trusted timestamping authority (RFC 3161), "
        "timestamps prove behavioral timing but not absolute temporal attestation. Whether this "
        "distinction matters in actual litigation is an open question for evidence law scholars."
    )

    doc.add_heading("2.5 Deviation Documentation Requirements", level=2)
    doc.add_paragraph(
        "When a radiologist's final assessment differs from their initial impression (especially "
        "after AI consultation), the platform enforces deviation documentation through the "
        "ClinicalReasoningDocumentor component. The component implements a 4-step wizard: "
        "(1) Acknowledge the AI finding, (2) Select clinical rationale from 15 pre-defined "
        "reason codes across 5 categories (anatomical, temporal, clinical, technical, "
        "interpretive), (3) Provide supporting evidence in free text, (4) Select follow-up "
        "recommendation. The protocol supports three enforcement levels: 'required' (must "
        "document), 'optional_with_attestation' (can skip with explicit attestation logged), "
        "and 'none' (no enforcement)."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 3: TECHNICAL ARCHITECTURE
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("3. Technical Architecture", level=1)
    doc.add_paragraph(
        "This section describes the implemented architecture as it exists in the codebase. "
        "Every component described here has corresponding source code. Where behavior is "
        "demo-quality rather than production-ready, this is noted."
    )

    doc.add_heading("3a. Event System", level=2)
    doc.add_paragraph(
        "The event system is defined across event_logger.ts and ExportPack.ts. Events follow "
        "a canonical format with typed payloads."
    )

    doc.add_heading("Event Types", level=3)

    # Event types table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ["Category", "Event Types", "Purpose"], header=True)
    table.rows[0].cells[0].text = ""

    event_cats = [
        ("Session", "SESSION_STARTED, RANDOMIZATION_ASSIGNED, SESSION_ENDED, EXPORT_GENERATED", "Session lifecycle and condition assignment"),
        ("Case", "CASE_LOADED, CASE_COMPLETED, IMAGE_VIEWED, VIEWPORT_CHANGED", "Case presentation and interaction"),
        ("Assessment", "FIRST_IMPRESSION_LOCKED, FINAL_ASSESSMENT, READ_EPISODE_STARTED/ENDED", "Core judgment recording with timing"),
        ("AI Consultation", "AI_REVEALED, DISCLOSURE_PRESENTED, DISCLOSURE_COMPREHENSION_RESPONSE", "AI output and error rate disclosure"),
        ("Deviation", "DEVIATION_STARTED, DEVIATION_SUBMITTED, DEVIATION_SKIPPED", "Override documentation or attestation"),
        ("Calibration", "CALIBRATION_STARTED, CALIBRATION_RESPONSE, CALIBRATION_FEEDBACK_SHOWN", "Pre-study calibration trials"),
        ("Attention", "ATTENTION_CHECK_PRESENTED, ATTENTION_CHECK_RESPONSE", "Validity checks during session"),
        ("Eye-tracking Proxy", "GAZE_ENTERED_ROI, GAZE_EXITED_ROI, DWELL_TIME_ROI, ATTENTION_COVERAGE_PROXY, SACCADE_DETECTED", "Mouse-based attention approximation (P1-3)"),
        ("Viewer", "VIEW_FOCUSED, ZOOM_CHANGED, PAN_CHANGED, WINDOW_LEVEL_CHANGED, AI_OVERLAY_TOGGLED", "Viewer interaction logging"),
        ("Workload", "WORKLOAD_THRESHOLD_CROSSED, WORKLOAD_ADVISORY_SHOWN/RESPONSE, SESSION_WORKLOAD_SUMMARY", "Fatigue and throughput monitoring"),
    ]

    for cat in event_cats:
        row = table.add_row()
        for i, val in enumerate(cat):
            row.cells[i].text = val
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.name = 'Calibri'

    doc.add_paragraph("")
    doc.add_heading("Event Sequence for HUMAN_FIRST Condition", level=3)
    doc.add_paragraph(
        "Based on the demo data (contrastDemoData.ts, hashChainDemoData.ts), a typical "
        "HUMAN_FIRST trial proceeds as follows:"
    )

    steps = [
        "CASE_LOADED — Case presented with condition assignment (HUMAN_FIRST)",
        "READ_EPISODE_STARTED (PRE_AI) — Independent reading begins; AI is not visible",
        "GAZE_ENTERED_ROI / DWELL_TIME_ROI — Attention proxy events logged as reader examines regions",
        "ATTENTION_COVERAGE_PROXY — Coverage summary before impression lock",
        "FIRST_IMPRESSION_LOCKED — Reader locks BI-RADS assessment with confidence rating; this is cryptographically committed before AI is visible",
        "READ_EPISODE_ENDED (PRE_AI) — Pre-AI reading period ends",
        "AI_REVEALED — AI output presented (suggested BI-RADS, confidence, localizations)",
        "DISCLOSURE_PRESENTED — FDR/FOR error rate disclosure shown to reader",
        "READ_EPISODE_STARTED (POST_AI) — Post-AI deliberation begins",
        "READ_EPISODE_ENDED (POST_AI) — Post-AI period ends",
        "FINAL_ASSESSMENT — Final BI-RADS with change analysis (TOWARD_AI / AWAY_FROM_AI / NO_CHANGE)",
        "TRUST_CALIBRATION — Pre-AI expectation vs. actual reliance self-report",
        "CASE_COMPLETED — Case concluded with total duration",
    ]

    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. ")
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run = p.add_run(step)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    doc.add_heading("Key Payload Structures", level=3)
    doc.add_paragraph(
        "FIRST_IMPRESSION_LOCKED payload: caseId, birads (0-6), confidence (1-5), "
        "timeToLockMs, viewerInteractions (zoomCount, panCount, windowLevelCount, viewsFocused[])."
    )
    doc.add_paragraph(
        "AI_REVEALED payload: caseId, suggestedBirads, aiConfidence (0-1), finding description, "
        "displayMode (NONE/GLOBAL_ONLY/BOX/CONTOUR/HEATMAP/SCALED_MARK)."
    )
    doc.add_paragraph(
        "FINAL_ASSESSMENT payload: caseId, birads, confidence, changeFromInitial (bool), "
        "changeDirection (TOWARD_AI/AWAY_FROM_AI/NO_CHANGE), postAiTimeMs."
    )

    p = doc.add_paragraph()
    run = p.add_run("What's NOT captured (gaps the panel should probe): ")
    run.bold = True
    run.font.name = 'Calibri'
    run = p.add_run(
        "Actual eye-tracking data (mouse proxy only); cognitive load measures beyond "
        "NASA-TLX self-report; dictation or verbal reasoning; multi-monitor viewport "
        "positions; radiologist identity verification at time of assessment."
    )

    doc.add_page_break()
    doc.add_heading("3b. Hash Chain (CanonicalHash.ts)", level=2)
    doc.add_paragraph(
        "Source: frontend/src/lib/CanonicalHash.ts (396 lines) and the browser-compatible "
        "implementation in ExportPack.ts (844 lines)."
    )

    doc.add_heading("Chain Construction Algorithm", level=3)
    p = doc.add_paragraph()
    run = p.add_run("Step 1 — Content Hash: ")
    run.bold = True
    p.add_run(
        "The event payload is serialized using RFC 8785 canonical JSON (sorted keys, no "
        "whitespace, normalized numbers, UTF-8). SHA-256 is computed over this canonical string. "
        "The canonicalJSON() function handles key ordering via Unicode code point comparison, "
        "number normalization (-0 → 0, NaN/Infinity → null), and minimal string escaping."
    )

    p = doc.add_paragraph()
    run = p.add_run("Step 2 — Chain Hash (128-byte structured input): ")
    run.bold = True
    p.add_run("A fixed-size buffer is constructed:")

    struct_items = [
        "Bytes 0-3: Sequence number (uint32, big-endian)",
        "Bytes 4-35: Previous chain hash (32 bytes, raw from hex)",
        "Bytes 36-71: Event ID (36 bytes, UUID as UTF-8, null-padded)",
        "Bytes 72-95: Timestamp (24 bytes, ISO 8601 as UTF-8, null-padded)",
        "Bytes 96-127: Content hash (32 bytes, raw from hex)",
    ]
    for item in struct_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item).font.size = Pt(10)

    doc.add_paragraph(
        "SHA-256 is computed over the entire 128-byte buffer to produce the chainHash. "
        "This structured encoding eliminates delimiter ambiguity — a critical fix "
        "identified in the Grayson Baird security review (noted in code comments as 'P0')."
    )

    p = doc.add_paragraph()
    run = p.add_run("Step 3 — Genesis Block: ")
    run.bold = True
    p.add_run(
        "The first entry uses GENESIS_HASH = '0'.repeat(64) (32 bytes of zeros as hex) "
        "for the previousHash field. This is a well-known constant, not a secret."
    )

    doc.add_heading("Verification", level=3)
    doc.add_paragraph(
        "The verifyChainFromData() function accepts raw events and ledger entries and "
        "performs the following checks: (1) Event count matches ledger count, "
        "(2) Sequence numbers are continuous 0..n-1, (3) Event IDs match between events "
        "and ledger, (4) Previous hash of entry i matches chain hash of entry i-1 (or "
        "GENESIS for i=0), (5) Content hash recomputed from payload matches recorded "
        "content hash, (6) Chain hash recomputed from structured input matches recorded "
        "chain hash. Failure returns the specific entry index and error type "
        "(SEQ_MISMATCH, EVENT_ID_MISMATCH, PREV_HASH_MISMATCH, CONTENT_TAMPERED, CHAIN_BROKEN)."
    )

    doc.add_heading("Trust Model Assessment", level=3)
    p = doc.add_paragraph()
    run = p.add_run("HONEST ASSESSMENT — What a cryptographer would challenge: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
    run.font.name = 'Calibri'

    doc.add_paragraph(
        "This entire system runs client-side. The hash chain proves internal consistency "
        "(no entry was modified after the fact), but it cannot prove when entries were "
        "created in absolute time. A sophisticated adversary with access to the client "
        "could construct a valid chain with fabricated timestamps. The manifest explicitly "
        "declares timestampTrustModel: 'client_clock_untrusted' to acknowledge this. "
        "Mitigations under development include RFC 3161 Timestamping Authority (TSA) "
        "integration (noted as P0 priority), server-side receipt generation "
        "(ServerReceipts.ts, 280 lines — implemented but not integrated), and Ed25519 "
        "digital signatures (referenced in ResearchDemoFlow.tsx but not yet operational). "
        "Until TSA or server receipts are integrated, the hash chain is tamper-evident "
        "but not time-attesting."
    )

    doc.add_heading("3c. Impression Ledger (ImpressionLedger.tsx)", level=2)
    doc.add_paragraph(
        "Source: frontend/src/components/research/legal/ImpressionLedger.tsx (712 lines). "
        "This is the legal-facing abstraction layer over the raw event stream."
    )

    doc.add_heading("Phase Model", level=3)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    add_table_row(table, ["Phase", "Entry Type", "What's Recorded"], header=True)
    table.rows[0].cells[0].text = ""

    phases = [
        ("first_impression", "HUMAN_FIRST_IMPRESSION", "BIRADSAssessment (category 0-6, confidence 1-5), aiVisible=false"),
        ("ai_exposure", "AI_OUTPUT_EXPOSURE", "AIFinding[] (id, score, flagged, region), DisclosureExposure (format, FDR/FOR metrics)"),
        ("(sub-phase)", "AI_FINDING_ACKNOWLEDGED", "Per-finding acknowledgment with timestamp and reviewed flag"),
        ("reconciliation", "RECONCILIATION", "Final BIRADSAssessment, optional DeviationRationale (reasonCodes[], supportingEvidence, followup)"),
        ("complete", "(terminal)", "Ledger sealed — no further entries accepted"),
    ]

    for phase in phases:
        row = table.add_row()
        for i, val in enumerate(phase):
            row.cells[i].text = val
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.name = 'Calibri'

    doc.add_paragraph("")
    doc.add_paragraph(
        "Each LedgerEntry contains: entryId, entryType, sequenceNumber, timestamp, "
        "timeOnTaskMs, assessment data, AI findings, disclosure config, acknowledgements, "
        "deviation rationale, caseId, readerId, previousHash, hash, and a locked flag. "
        "The hash chain within the ledger uses Web Crypto API (crypto.subtle.digest) "
        "for SHA-256 computation, making it browser-native without external dependencies."
    )

    doc.add_paragraph(
        "The temporal ordering proof works as follows: the HUMAN_FIRST_IMPRESSION entry's "
        "hash is incorporated into the AI_OUTPUT_EXPOSURE entry's previousHash field. "
        "If the first impression were altered after AI exposure, the chain link between "
        "entries 1 and 2 would fail verification. This is the core legal artifact — the "
        "thing that survives deposition cross-examination."
    )

    doc.add_heading("3d. Export Pack", level=2)
    doc.add_paragraph(
        "Source: frontend/src/lib/ExportPack.ts (844 lines) and ExportPackZip.ts (1046 lines). "
        "A complete export contains 6 files:"
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    add_table_row(table, ["File", "Format", "Contents"], header=True)
    table.rows[0].cells[0].text = ""

    export_files = [
        ("trial_manifest.json", "JSON", "Session metadata, integrity checksums, protocol config, timestamp trust model, file checksums, disclosure provenance, randomization metadata"),
        ("events.jsonl", "JSONL", "Append-only event stream, one JSON object per line, sequenced 0..n-1"),
        ("ledger.json", "JSON", "Hash chain entries with contentHash, previousHash, chainHash for each event"),
        ("verifier_output.json", "JSON", "Automated integrity check results: 8 checks including chain integrity, timestamp monotonicity, pre-AI lock verification, disclosure logging"),
        ("derived_metrics.csv", "CSV", "Pre-computed analysis variables — one row per session/case"),
        ("codebook.md", "Markdown", "Field definitions, event type documentation, ADDA operational definition, hash chain specification"),
    ]

    for f in export_files:
        row = table.add_row()
        for i, val in enumerate(f):
            row.cells[i].text = val
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.name = 'Calibri'

    doc.add_paragraph("")
    doc.add_paragraph(
        "The TrialManifest includes: exportVersion, schemaVersion, exportTimestamp, sessionId, "
        "participantId, condition, integrity block (eventCount, finalHash, chainValid), "
        "protocol block (revealTiming, disclosureFormat, deviationEnforcement), "
        "timestampTrustModel, fileChecksums (SHA-256 of each file), disclosureProvenance "
        "(FDR/FOR values and their source), and randomization metadata (seed, assignment "
        "method, condition matrix hash). In litigation, this package would be presented as "
        "a self-contained, independently verifiable record of the clinical decision process."
    )

    doc.add_heading("3e. Derived Metrics", level=2)
    doc.add_paragraph(
        "Source: DerivedMetrics interface in ExportPack.ts. These are the pre-computed "
        "analysis variables exported with each session."
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    add_table_row(table, ["Metric", "Type", "Clinical/Legal Significance"], header=True)
    table.rows[0].cells[0].text = ""

    metrics = [
        ("initialBirads / finalBirads", "int", "Assessment values before and after AI — core of change analysis"),
        ("aiBirads / aiConfidence", "int/float", "AI recommendation for comparison"),
        ("changeOccurred", "bool", "Whether assessment changed post-AI"),
        ("aiConsistentChange", "bool", "Changed AND final matches AI — possible automation bias indicator"),
        ("aiInconsistentChange", "bool", "Changed AND final differs from AI — independent clinical judgment"),
        ("adda", "bool|null", "ADDA (Automation Decision-Delegation Assessment): TRUE if changed toward AI when initial disagreed with AI. NULL if not in denominator (initial already matched AI)"),
        ("addaDenominator", "bool", "TRUE if initial ≠ AI — eligible for ADDA calculation"),
        ("deviationRequired / Documented / Skipped", "bool", "Deviation documentation status"),
        ("preAiReadMs", "int", "Time spent on independent read before AI — key defensibility metric"),
        ("postAiReadMs", "int", "Time spent after AI consultation"),
        ("timeRatio", "float", "preAiReadMs / postAiReadMs — higher = more pre-AI diligence"),
        ("sessionMedianPreAITime", "int", "Baseline for within-session comparison (Hasty Review Index)"),
        ("preAITimeVsMedian", "int", "Deviation from session median — flags rushing"),
        ("lockToRevealMs", "int", "Time from impression lock to AI reveal"),
        ("revealToFinalMs", "int", "Time from AI reveal to final assessment"),
        ("decisionChangeCount", "int", "Number of assessment changes in session"),
        ("comprehensionCorrect", "bool|null", "FDR/FOR comprehension check result"),
        ("timingFlagPreAiTooFast", "bool", "Pre-AI read < 3s threshold"),
        ("timingFlagAiExposureTooFast", "bool", "AI exposure < 3s threshold"),
    ]

    for m in metrics:
        row = table.add_row()
        for i, val in enumerate(m):
            row.cells[i].text = val
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.name = 'Calibri'

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("FOR QUANTITATIVE PSYCHOLOGISTS: ")
    run.bold = True
    run.font.name = 'Calibri'
    run = p.add_run(
        "These metrics are operationally defined in code but have not been psychometrically "
        "validated. The ADDA metric has a clear operational definition (documented in the "
        "codebook) but no published validation study. The timing thresholds (3s for 'too fast,' "
        "15s for adequate pre-AI time) are derived from research literature and internal "
        "judgment, not from empirical calibration against gold-standard performance measures. "
        "Test-retest reliability, construct validity, and sensitivity to meaningful clinical "
        "differences are unknown."
    )

    doc.add_heading("3f. Local-First Architecture", level=2)
    doc.add_paragraph(
        "Evidify is designed as a local-first application: no Protected Health Information (PHI) "
        "leaves the device. The Rust/Tauri backend provides SQLCipher-encrypted storage "
        "(AES-256), Argon2id key derivation (m=64MB, t=3, p=4), and OS keychain integration "
        "for wrapped key storage. The frontend runs as a Tauri webview with Content Security "
        "Policy restrictions."
    )

    p = doc.add_paragraph()
    run = p.add_run("FIPS compliance: ")
    run.bold = True
    p.add_run(
        "Noted as P0 priority in the codebase but not yet implemented. The current "
        "implementation uses standard Web Crypto API and Node.js crypto, which are not "
        "FIPS-validated modules."
    )

    p = doc.add_paragraph()
    run = p.add_run("TSA (Timestamping Authority) integration: ")
    run.bold = True
    p.add_run(
        "ServerReceipts.ts (280 lines) implements a server-side receipt system for "
        "tamper-evident timestamps with batch processing and verification. This is "
        "implemented as code but not integrated into the live workflow. RFC 3161 "
        "compliance is referenced but not achieved."
    )

    p = doc.add_paragraph()
    run = p.add_run("FOR CRYPTOGRAPHERS/SECURITY: ")
    run.bold = True
    run.font.name = 'Calibri'
    run = p.add_run(
        "The trust model boundary is the client device. Attack surfaces include: "
        "(1) Client-side timestamp manipulation before chain construction, "
        "(2) Browser extension injection into the web crypto context, "
        "(3) Tauri IPC interception between frontend and Rust backend, "
        "(4) SQLCipher key extraction from OS keychain, "
        "(5) No certificate pinning for any future network calls. "
        "The README's security claims table is refreshingly honest about what's "
        "verifiable vs. assumed."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 4: THE CONTRAST — STI PREVENTION
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("4. The Contrast — STI Prevention Thesis", level=1)

    doc.add_paragraph(
        "Source: frontend/src/components/research/legal/TheContrast.tsx (613 lines). "
        "This component is the conceptual centerpiece of Evidify's legal strategy."
    )

    doc.add_heading("4.1 Theoretical Framework", level=2)
    doc.add_paragraph(
        "The Contrast builds on two established bodies of research:"
    )

    p = doc.add_paragraph()
    run = p.add_run("The Story Model (Pennington & Hastie, 1992): ")
    run.bold = True
    p.add_run(
        "Jurors do not evaluate evidence analytically. They construct coherent narrative "
        "accounts and select the verdict that best fits their constructed story. Evidence "
        "that doesn't fit into a story is underweighted or ignored. A bare timestamp log "
        "provides no story — jurors must construct their own, and plaintiff's counsel will "
        "help them construct an unfavorable one."
    )

    p = doc.add_paragraph()
    run = p.add_run("Spontaneous Trait Inference (Uleman et al., 2008): ")
    run.bold = True
    p.add_run(
        "Upon encountering behavioral descriptions, observers automatically and "
        "involuntarily infer dispositional traits about the actor. '2 minutes, 14 seconds' "
        "triggers the inference 'careless.' This is not a deliberate judgment — it occurs "
        "below conscious awareness and is resistant to correction once formed. The inference "
        "operates on the behavioral description itself, not on any explicit evaluative "
        "statement."
    )

    doc.add_heading("4.2 The Two-Panel Demonstration", level=2)
    doc.add_paragraph(
        "TheContrast.tsx renders two panels side by side, using identical underlying data "
        "from the trial event stream and derived metrics:"
    )

    p = doc.add_paragraph()
    run = p.add_run("Left Panel — 'Prosecution Exhibit' (StructuredLogPanel): ")
    run.bold = True
    p.add_run(
        "A monospace table showing: Event | Timestamp | Duration. Rows include 'Case loaded,' "
        "'First impression locked,' 'AI consultation initiated,' 'Final assessment submitted,' "
        "and 'Total session.' Timing is displayed as raw values (e.g., 'Pre-AI read time: "
        "18 sec'). The footer shows total time in bold red. This presentation strips all "
        "clinical context and invites STI."
    )

    p = doc.add_paragraph()
    run = p.add_run("Right Panel — 'Defense Exhibit' (ClinicalNarrativePanel): ")
    run.bold = True
    p.add_run(
        "A serif-font narrative paragraph that contextualizes the same data: 'The reader "
        "began independent review of the bilateral mammographic study at [time]. During the "
        "[duration] pre-consultation period, the reader assessed [density] breast tissue, "
        "examined [N] standard views, and formed an initial clinical impression of [BI-RADS] "
        "with a confidence level of [X/5]. This independent assessment was cryptographically "
        "locked before any AI output was visible to the reader.' The same temporal data is "
        "present but embedded in a clinical reasoning framework that disrupts STI."
    )

    doc.add_paragraph(
        "The component includes a 'Highlight Shared Data' toggle that applies a yellow "
        "highlight to identical data points across both panels, making it visually obvious "
        "that the factual content is unchanged. A 'Swap Panels' button allows viewers to "
        "see each framing in the other position."
    )

    doc.add_heading("4.3 Research Context", level=2)
    doc.add_paragraph(
        "The ResearchContextSection (collapsible) explains the STI mechanism to researchers "
        "and provides citations. This was built specifically for the February 13, 2026 meeting "
        "with Brown University researchers studying how jurors perceive radiologist behavior "
        "in AI-assisted decision-making contexts."
    )

    p = doc.add_paragraph()
    run = p.add_run("FOR BEHAVIORAL ECONOMISTS / HCI RESEARCHERS: ")
    run.bold = True
    run.font.name = 'Calibri'
    run = p.add_run(
        "The STI prevention claim is theoretically grounded but experimentally unvalidated "
        "in this specific application. Key questions: (1) Does the narrative framing actually "
        "reduce unfavorable trait inferences about the radiologist in a mock jury study? "
        "(2) Does it reduce them enough to affect verdict preferences? (3) Are there "
        "confounds — does the narrative introduce its own biases (e.g., the warm glow of "
        "professional language)? (4) Does the highlighting feature help or hurt — does "
        "drawing attention to shared data make the framing effect more or less salient? "
        "(5) What is the effect of panel ordering (prosecution first vs. defense first)? "
        "These are precisely the questions the Brown collaboration is designed to answer."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 5: CLINICAL WORKFLOW INTEGRATION
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("5. Clinical Workflow Integration", level=1)

    doc.add_paragraph(
        "Source: frontend/src/components/ResearchDemoFlow.tsx (7,424 lines). This is the "
        "primary research demonstration component implementing the full mammography reading "
        "workflow."
    )

    doc.add_heading("5.1 Workflow Phases", level=2)

    phases_desc = [
        ("Setup", "Study configuration: protocol selection (BRPLL-MAMMO-v1.0), condition assignment via Latin Square 4×4 counterbalancing, randomization seed, disclosure format (FDR/FOR), case queue generation (1 calibration + 3 study cases)."),
        ("Calibration", "Pre-study calibration trial with known ground truth. Reader assesses a case and receives feedback on accuracy. This establishes baseline behavior and familiarizes the reader with the interface. Calibration cases are marked with isCalibration flag and excluded from study analyses."),
        ("Independent Reading (Pre-AI)", "Reader views mammogram (4 standard views: LCC, LMLO, RCC, RMLO), interacts with the viewer (zoom, pan, window/level), and forms initial BI-RADS assessment. AI output is not visible. Keyboard shortcuts: 0-5 for BI-RADS, Space to lock impression, confidence +/- keys. Viewport attention tracking logs which anatomical regions were examined and for how long."),
        ("AI Consultation with Disclosure", "After locking the first impression, AI output is revealed: suggested BI-RADS, confidence score, lesion localizations (bounding boxes). Error rate disclosure is presented in the configured format — FDR 4%, FOR 12% in the demo data. A comprehension check asks the reader to demonstrate understanding of the error rates."),
        ("Post-AI Reconciliation", "Reader reviews AI output, examines flagged regions, and submits final assessment. If the final assessment differs from the initial impression, the deviation documentation workflow is triggered: acknowledge AI finding → select reason codes → provide evidence → specify follow-up."),
        ("Export and Verification", "Session data is packaged into the export format. Hash chain integrity is verified. Expert witness packet is generated. ZIP file is downloadable."),
    ]

    for name, desc in phases_desc:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        run.font.name = 'Calibri'
        p.add_run(desc)

    doc.add_heading("5.2 Study States", level=2)
    doc.add_paragraph(
        "The demo flow tracks state as: SETUP → CALIBRATION → CALIBRATION_FEEDBACK → "
        "INITIAL → AI_REVEALED → DEVIATION → COMPLETE → STUDY_COMPLETE. Session recovery "
        "is supported via localStorage with a 24-hour window."
    )

    doc.add_heading("5.3 Supplementary Instruments", level=2)
    doc.add_paragraph(
        "The ResearchDemoFlow includes several embedded research instruments: "
        "(1) NASA-TLX workload assessment (ProbesBatchModal), "
        "(2) Trust calibration probes (pre-AI expectation vs. actual reliance), "
        "(3) Comprehension checks on error rate interpretation, "
        "(4) Attention checks for session validity, "
        "(5) AutomationBiasRiskMeter (weighted risk from lock time, agreement streak, "
        "deviation skip rate, AI fixation dwell), "
        "(6) ValidityResponseStylePanel (MMPI-inspired HRI, CPI, DAI, ENG indicators), "
        "(7) LiabilityRiskPanel (Baird framework: LOW/MODERATE/HIGH/CRITICAL classification), "
        "(8) CrossExaminationPanel (vulnerability analysis and defense strengths)."
    )

    p = doc.add_paragraph()
    run = p.add_run("FOR RADIOLOGISTS: ")
    run.bold = True
    run.font.name = 'Calibri'
    run = p.add_run(
        "This workflow adds steps that do not exist in current clinical practice. The "
        "independent impression lock, explicit AI acknowledgment, and deviation documentation "
        "all represent additional friction. In a production environment, this friction must "
        "be justified by the legal and quality-improvement value it provides. The keyboard "
        "shortcut system (0-5 for BI-RADS, Space to lock) attempts to minimize the additional "
        "time cost, but real-world validation with practicing radiologists has not been conducted."
    )

    p = doc.add_paragraph()
    run = p.add_run("FOR MEDICAL EDUCATION: ")
    run.bold = True
    run.font.name = 'Calibri'
    run = p.add_run(
        "The instrumentation captures decision-making patterns that could be valuable for "
        "training. The hash-chained event log documents how a trainee's reading behavior "
        "changes in response to AI input — whether they develop appropriate reliance or "
        "develop automation bias over time. The validity indicators (HRI, CPI, DAI, ENG) "
        "could flag concerning patterns for supervisor review. The calibration trial provides "
        "a baseline assessment capability."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 6: WORKLOAD MONITORING
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("6. Workload Monitoring", level=1)

    doc.add_paragraph(
        "Source: frontend/src/components/research/WorkloadMonitor.tsx (218 lines) and "
        "frontend/src/types/workloadTypes.ts (comprehensive type definitions)."
    )

    doc.add_heading("6.1 Tracked Metrics", level=2)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    add_table_row(table, ["Metric", "Computation", "Purpose"], header=True)
    table.rows[0].cells[0].text = ""

    wl_metrics = [
        ("casesCompleted", "Count of completed cases in session", "Volume tracking"),
        ("totalReadingTimeMs", "Cumulative reading time", "Fatigue assessment"),
        ("averageTimePerCaseMs", "totalReadingTime / casesCompleted", "Efficiency monitoring"),
        ("casesPerHour", "Real-time throughput rate", "Workload intensity"),
        ("workloadStatus", "GREEN (<30/hr) / YELLOW (30-40/hr) / RED (>40/hr)", "Threshold-based alerting"),
        ("sessionDurationIndex", "0-100 composite of duration, cases, and rate", "Overall session intensity"),
    ]

    for m in wl_metrics:
        row = table.add_row()
        for i, val in enumerate(m):
            row.cells[i].text = val
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.name = 'Calibri'

    doc.add_heading("6.2 Status Indicators and Thresholds", level=2)
    doc.add_paragraph(
        "Default thresholds from DEFAULT_WORKLOAD_THRESHOLDS: casesPerHourYellow = 30, "
        "casesPerHourRed = 40, maxSessionCases = 50, maxSessionMinutes = 180. The "
        "sessionDurationIndex bar uses green (<40%), amber (40-70%), red (>70%) "
        "color coding. Workload advisory events are logged when thresholds are crossed, "
        "and user responses (CONTINUE / TAKE_BREAK) are recorded."
    )

    doc.add_heading("6.3 Cohort Comparison", level=2)
    doc.add_paragraph(
        "The WorkloadDashboard component (separate from WorkloadMonitor) provides cohort "
        "comparison against practice medians. Demo data shows a 14-radiologist practice "
        "with median 26.1 cases/hour and mean 2:12/case. Individual performance is compared "
        "against practice P25 and P75."
    )

    p = doc.add_paragraph()
    run = p.add_run("FOR HUMAN FACTORS EXPERTS: ")
    run.bold = True
    run.font.name = 'Calibri'
    run = p.add_run(
        "The threshold values (30/hr yellow, 40/hr red) are referenced to radiologist "
        "workload research showing performance degradation beyond certain throughput "
        "levels, but specific citations are not included in the code. The sessionDurationIndex "
        "is a composite metric without published derivation or validation. The 180-minute "
        "maximum session duration is a reasonable default but should be validated against "
        "actual fatigue research in mammography. These thresholds are configurable "
        "(WorkloadThresholds interface), allowing institutional customization, but the "
        "question of what values are clinically meaningful remains open."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 7: HONEST GAPS
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("7. What's Not Built Yet — Honest Gaps", level=1)

    doc.add_paragraph(
        "This section is deliberately unflattering. Its purpose is to give the panel "
        "an accurate picture of what exists versus what is aspirational."
    )

    # Gap analysis table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    add_table_row(table, ["Area", "Status", "What Exists", "What's Missing"], header=True)
    table.rows[0].cells[0].text = ""

    gaps = [
        ("Eye Tracking", "Schema only",
         "GAZE_ENTERED_ROI and DWELL_TIME_ROI events exist in the event schema. EyeTrackingHooks.ts (461 lines) implements WebGazer and Tobii integration with fixation detection and scan path analysis.",
         "No actual eye tracker is connected. All 'gaze' events in the demo are mouse-hover proxies. The EyeTrackingHooks code is implemented but untested with real hardware."),
        ("FIPS Compliance", "Not started",
         "Referenced as P0 priority in code comments.",
         "No FIPS-validated crypto modules are used. Web Crypto API and Node.js crypto are not FIPS-certified. Would require FIPS 140-2/3 validated implementations."),
        ("TSA Integration", "Code exists, not wired",
         "ServerReceipts.ts implements server-side receipt system. RFC 3161 referenced.",
         "No connection to any timestamping authority. No server infrastructure exists. Timestamps remain client_clock_untrusted."),
        ("Real Mammogram Images", "Placeholders only",
         "INbreast case references exist (caseId '20586908'). Image dimensions are defined (3328×4096). Setup scripts exist for INbreast data linking.",
         "No actual DICOM or PNG images are loaded. webImagePath is undefined in all demo cases. The viewer renders placeholder regions."),
        ("EHR Integration", "None",
         "PACSIntegration.ts (500 lines) defines DICOM retrieval and worklist interfaces. Tauri IPC bindings exist for vault management.",
         "No actual PACS, RIS, or EHR connections. All interfaces are stubs. HL7 FHIR integration is not implemented."),
        ("Legal Review", "Not done",
         "FRE 902(13)/(14) and Vermont §1913 are referenced in code and component text.",
         "No attorney has reviewed the legal claims. The FRE 902 compliance assertion is an architectural goal, not a legal opinion."),
        ("Psychometric Validation", "Not done",
         "ADDA, HRI, CPI, DAI, ENG metrics are operationally defined. Thresholds are set.",
         "No validation study. No test-retest reliability data. No construct validity assessment. No sensitivity analysis."),
        ("IRB Approval", "Not obtained",
         "Study protocol components exist (StudyProtocol.tsx). Research design is documented.",
         "No IRB application has been submitted. No approval obtained from any institution."),
        ("HIPAA Compliance", "Architectural claim only",
         "Local-first design means no PHI leaves device. CSP restricts network access.",
         "No formal HIPAA audit. No BAA templates. No HIPAA compliance officer review. The claim is architectural, not certified."),
        ("Multi-site Deployment", "Not built",
         "SiteRoleConfig.ts (363 lines) defines multi-site roles, permissions, and IRB tracking.",
         "No deployment infrastructure. No site management console. No data aggregation across sites."),
        ("AI Research Assistant", "Not built",
         "Tauri IPC bindings reference AI analysis, deep analysis, and supervision commands.",
         "These are stub interfaces. No AI assistant functionality exists in the research instrumentation."),
        ("Production Hardening", "Not done",
         "Source maps disabled. IP protection noted. Tauri security architecture.",
         "No penetration testing. No security audit. No load testing. No accessibility audit. No localization."),
    ]

    for g in gaps:
        row = table.add_row()
        for i, val in enumerate(g):
            row.cells[i].text = val
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7.5)
                    run.font.name = 'Calibri'

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("FOR EVERY EXPERT: ")
    run.bold = True
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
    run = p.add_run(
        "Before trusting this platform in your domain, what would you need to see? The "
        "architect asks this question sincerely. The codebase demonstrates architectural "
        "intent and technical competence, but the gap between 'demo that works' and "
        "'system you'd stake your career on' is substantial. The purpose of this panel "
        "review is to map that gap precisely."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 8: RESEARCH DESIGN
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("8. Research Design — Proposed BRPLL Collaboration", level=1)

    doc.add_paragraph(
        "This section presents Henderson's proposed study design for the Brown "
        "Radiology-Psychology Legal Lab (BRPLL) collaboration. This is a proposed "
        "research protocol, not an approved study."
    )

    doc.add_heading("8.1 Study Population and Setting", level=2)
    doc.add_paragraph(
        "Participants: Board-certified radiologists and radiology residents with "
        "mammography reading experience. Setting: Controlled laboratory environment "
        "using the Evidify research platform. Cases: Mammography studies from the INbreast "
        "dataset (de-identified, publicly available research dataset) with known pathology "
        "ground truth."
    )

    doc.add_heading("8.2 Experimental Design", level=2)
    doc.add_paragraph(
        "The study uses a within-subjects multi-reader multi-case (MRMC) design "
        "with counterbalanced conditions:"
    )

    p = doc.add_paragraph()
    run = p.add_run("Conditions: ")
    run.bold = True
    p.add_run(
        "HUMAN_FIRST (reader forms independent impression before AI reveal) vs. "
        "AI_FIRST (AI output presented before reader assessment) vs. CONCURRENT "
        "(AI and reader assess simultaneously). Condition assignment uses Latin Square "
        "4×4 counterbalancing with stratified block randomization, implemented in "
        "condition_matrix.ts."
    )

    p = doc.add_paragraph()
    run = p.add_run("Disclosure Manipulation: ")
    run.bold = True
    p.add_run(
        "FDR/FOR format (numeric error rates) vs. natural frequency format vs. no "
        "disclosure. The demo uses FDR 4% and FOR 12% with a calibration sample of "
        "N=2,500 at 5% prevalence."
    )

    doc.add_heading("8.3 Primary Outcomes", level=2)

    outcomes = [
        ("ADDA Rate", "Proportion of cases where reader changed toward AI recommendation when initial assessment disagreed. This is the primary measure of automation bias."),
        ("Decision Change Patterns", "Direction (TOWARD_AI vs. AWAY_FROM_AI), magnitude (BI-RADS category distance), and timing of assessment changes."),
        ("Timing Metrics", "Pre-AI reading time, post-AI deliberation time, time ratio, and timing flag rates (pre-AI < 3s, AI exposure < 3s)."),
    ]

    for name, desc in outcomes:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        run.font.name = 'Calibri'
        p.add_run(desc)

    doc.add_heading("8.4 Secondary Outcomes", level=2)

    secondary = [
        ("Trust Calibration", "Pre-AI expectation of AI helpfulness vs. actual reliance (TRUST_CALIBRATION event)."),
        ("Workload Effects", "Session duration index, cases-per-hour trends, workload threshold crossings."),
        ("Documentation Quality", "Deviation rationale word count, reason code selection patterns, documentation avoidance index (DAI)."),
        ("Validity Indicators", "HRI, CPI, DAI, ENG profiles across conditions — do response patterns differ by condition?"),
        ("Comprehension", "FDR/FOR comprehension check accuracy — does error rate understanding mediate ADDA?"),
    ]

    for name, desc in secondary:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        run.font.name = 'Calibri'
        p.add_run(desc)

    doc.add_heading("8.5 The 'Iff' Condition", level=2)
    doc.add_paragraph(
        "The documentation proves independent reading if and only if: (1) the hash chain "
        "is intact (all verification checks pass), AND (2) the FIRST_IMPRESSION_LOCKED event "
        "precedes the AI_REVEALED event in the chain, AND (3) the chain link between these "
        "events is verified (the AI_REVEALED entry's previous hash matches the hash of the "
        "entry containing or following FIRST_IMPRESSION_LOCKED). If any of these conditions "
        "fail, the independence claim is falsified — and the falsification is machine-verifiable."
    )

    p = doc.add_paragraph()
    run = p.add_run("FOR BIOSTATISTICIANS: ")
    run.bold = True
    run.font.name = 'Calibri'
    run = p.add_run(
        "Power analysis assumptions have not been computed. Key considerations: "
        "(1) Effect size for ADDA under HUMAN_FIRST vs. AI_FIRST is unknown — prior "
        "literature on automation bias in radiology suggests moderate effects but varies "
        "widely. (2) The MRMC design requires modeling reader × case interaction, typically "
        "using Obuchowski-Rockette or Hillis methods. (3) iMRMC export format is supported "
        "(iMRMCExport.ts, 416 lines) with R script generation for FDA-standard MRMC analysis. "
        "(4) Sample size will be constrained by radiologist availability. "
        "(5) Multiple comparison corrections needed for the number of metrics examined."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 9: MARKET & STRATEGIC POSITION
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("9. Market & Strategic Position", level=1)

    doc.add_heading("9.1 Target Users", level=2)

    users = [
        ("Near-term — Forensic Psychologists", "The Tauri/vault architecture was originally designed for forensic psychology documentation. At $280/hour forensic rates, time savings from structured documentation and automated report generation represent clear ROI. The forensic workspace, testimony mode, and claim ledger components exist in the codebase."),
        ("Medium-term — Radiologists (via BRPLL)", "The research collaboration with Brown establishes the evidence base for clinical adoption. Mammography is the beachhead specialty due to high AI adoption, high liability exposure, and high-volume reading sessions."),
        ("Long-term — Broader Clinical AI Users", "Any clinician using AI-assisted decision tools faces the same liability documentation problem. The architecture is not specialty-specific — the event system, hash chain, and export format are generalizable."),
    ]

    for name, desc in users:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        run.font.name = 'Calibri'
        p.add_run(desc)

    doc.add_heading("9.2 Competitive Landscape", level=2)
    doc.add_paragraph(
        "Honest assessment: the codebase does not contain competitive analysis data. What "
        "can be observed from the architecture: no existing PACS vendor provides "
        "cryptographically chained decision documentation. EHR systems record what was "
        "decided but not the temporal relationship between human judgment and AI consultation. "
        "AI vendors (e.g., iCAD, Hologic, Lunit) provide their own performance metrics but "
        "not documentation of the human decision process. The gap Evidify targets — "
        "court-defensible documentation of the human-AI interaction — appears to be unoccupied, "
        "but this claim has not been rigorously validated through market research."
    )

    doc.add_heading("9.3 Business Model Implications", level=2)

    models = [
        "SaaS for forensic documentation (monthly subscription per clinician)",
        "Research licensing (per-site, per-study license for academic use)",
        "Expert witness support tools (per-case packaging and analysis)",
        "Integration licensing (PACS/EHR vendors embedding the documentation layer)",
    ]

    for m in models:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(m)

    p = doc.add_paragraph()
    run = p.add_run("FOR VCs: ")
    run.bold = True
    run.font.name = 'Calibri'
    run = p.add_run(
        "TAM estimation requires data the codebase doesn't contain. Rough framing: "
        "~35,000 radiologists in the US, growing AI adoption rates, rising malpractice "
        "premiums. The moat is (1) the hash chain specification becoming an accepted "
        "evidentiary standard, (2) validated psychometric instruments with published "
        "norms, and (3) published research from the BRPLL collaboration. What you'd need "
        "in a Series A deck: validated clinical data, at least one legal precedent or "
        "judicial recognition of the documentation standard, letters of intent from "
        "radiology practices, and a PACS integration partnership."
    )

    p = doc.add_paragraph()
    run = p.add_run("FOR HEALTH SYSTEM INNOVATION: ")
    run.bold = True
    run.font.name = 'Calibri'
    run = p.add_run(
        "Procurement would require: SOC 2 Type II certification (the README explicitly "
        "states 'not certified'), HIPAA BAA, HL7 FHIR interoperability, PACS integration, "
        "IT security review, and clinical workflow validation study. The local-first "
        "architecture simplifies some procurement concerns (no PHI in cloud) but complicates "
        "others (per-device deployment, no centralized management)."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 10: VISION
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("10. Vision — What This Could Become", level=1)

    doc.add_paragraph(
        "This section is intentionally forward-looking. It describes what the platform "
        "could become if the research validates the core thesis and adoption follows."
    )

    doc.add_heading("10.1 Autonomous Documentation Standard", level=2)
    doc.add_paragraph(
        "If the hash-chained temporal documentation approach is validated by the BRPLL "
        "research and accepted in legal proceedings, it could become the standard for "
        "documenting any AI-assisted clinical decision. The schema is not radiology-specific: "
        "the event types, hash chain, and derived metrics pattern applies to any workflow "
        "where a human makes a judgment, consults an AI, and produces a final decision."
    )

    doc.add_heading("10.2 Cross-Specialty Expansion", level=2)
    doc.add_paragraph(
        "Natural expansion path: radiology (mammography → chest CT → other modalities) → "
        "pathology (digital pathology with AI-assisted grading) → dermatology (dermoscopy "
        "with AI triage) → any image-based diagnosis. The viewport attention tracking "
        "framework (viewportAttentionTypes.ts) generalizes to any image examination workflow. "
        "The error classification taxonomy (wolfeClassification.ts — search, recognition, "
        "and decision errors based on visual attention research) is applicable across "
        "image-interpretation specialties."
    )

    doc.add_heading("10.3 Regulatory Integration", level=2)
    doc.add_paragraph(
        "Could this become what CMS or FDA requires? The FDA's approach to AI/ML-based "
        "Software as a Medical Device (SaMD) focuses on algorithm performance. There is "
        "currently no regulatory requirement for documenting the human-AI interaction "
        "process. If malpractice litigation drives demand for this documentation, regulatory "
        "requirements could follow. The export format (trial_manifest.json + events.jsonl + "
        "ledger.json) is designed to be machine-readable by regulatory reviewers."
    )

    doc.add_heading("10.4 Training and Credentialing", level=2)
    doc.add_paragraph(
        "Hash-chained decision logs could become part of board certification or "
        "continuing education: demonstrating that a radiologist develops appropriate "
        "AI reliance patterns (not rubber-stamping, not ignoring) over time. The validity "
        "indicators (HRI, CPI, DAI, ENG) could establish normative ranges for 'healthy' "
        "AI interaction patterns. The calibration trial framework already supports "
        "before/after assessment of reading behavior."
    )

    doc.add_heading("10.5 Insurance Impact", level=2)
    doc.add_paragraph(
        "If verified autonomous decision-making documentation reduces malpractice exposure, "
        "insurers could offer premium reductions for practices that adopt the documentation "
        "standard. The liability classifier (liabilityClassifier.ts) already categorizes "
        "cases into risk levels (LOW/MODERATE/HIGH/CRITICAL) based on the relationship "
        "between initial assessment, AI recommendation, and final assessment. Aggregate "
        "risk profiles could feed actuarial models."
    )

    doc.add_heading("10.6 Legal Precedent", level=2)
    doc.add_paragraph(
        "If this documentation standard is adopted and a case is litigated where the "
        "hash chain evidence is admitted under FRE 902(13)/(14), it creates precedent "
        "that shapes subsequent cases. The first successful admission of Evidify-format "
        "evidence would be a significant inflection point for adoption."
    )

    p = doc.add_paragraph()
    run = p.add_run("FOR THE FULL PANEL: ")
    run.bold = True
    run.font.name = 'Calibri'
    run = p.add_run(
        "What is the most important thing this platform could become that the architect "
        "hasn't considered? What adjacent application would create the most value? What "
        "would make you, in your specific domain, want to use or recommend this?"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 11: QUESTIONS FOR THE PANEL
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("11. Questions for the Panel", level=1)

    doc.add_paragraph(
        "The following questions are designed to be specific and actionable. They target "
        "the areas where expert input would most change the platform's development trajectory."
    )

    panel_questions = {
        "For Diagnostic Radiologists": [
            "What is the minimum pre-AI reading time (in seconds) you would consider adequate for a screening mammogram with heterogeneously dense tissue? Does the 15-second threshold in our rubber-stamp detection align with clinical practice?",
            "Would you use this tool if it added 30 seconds per case to your workflow? 60 seconds? What's the threshold where the documentation value no longer justifies the time cost?",
            "Is the BI-RADS 0-6 + confidence 1-5 assessment structure sufficient for capturing your clinical impression, or do you need additional structured fields (e.g., lesion description, location)?",
            "How would you react if a malpractice attorney presented your hash-chained decision log as evidence — would you view it as protective or as creating additional attack surface?",
        ],
        "For Evidence Law Scholars": [
            "Does this hash chain implementation meet the technical requirements for self-authentication under FRE 902(13) or 902(14)? What specific deficiencies would you identify?",
            "Is the 'client_clock_untrusted' timestamp trust model a fatal flaw for evidentiary purposes, or is behavioral timing sufficient when combined with chain integrity?",
            "Would a court accept the tamper detection demonstration (showing what chain breakage looks like) as part of foundation testimony? Is this educational approach effective with juries?",
            "What Vermont-specific considerations arise under 12 V.S.A. §1913 that differ from federal requirements?",
        ],
        "For Cryptographers": [
            "Is the 128-byte structured encoding for chain hash input sufficiently unambiguous? Are there edge cases in the UUID padding or timestamp truncation that could create collisions?",
            "Given that this runs client-side, what is the minimum viable server-side component needed to make timestamp claims credible? Is RFC 3161 TSA sufficient, or is something stronger needed?",
            "The canonical JSON implementation omits RFC 8785 handling of Unicode normalization forms. Is this a practical concern for the types of payloads in this system?",
            "What would you need to verify before signing off on the hash chain as 'cryptographically sound for its stated purpose'?",
        ],
        "For Forensic Psychologists": [
            "Are the MMPI-inspired validity indicators (HRI, CPI, DAI, ENG) appropriate adaptations for this context, or do they overfit to the MMPI framework?",
            "The interpretive summary generation produces MMPI-style narrative text. Is this format appropriate for research contexts, or does it imply clinical validation that doesn't exist?",
            "Would you use this documentation format for forensic evaluations in your practice? What additional fields or instruments would you require?",
        ],
        "For Health Tech Venture Capital": [
            "What is the minimum evidence threshold for a Series A in this category — published research, legal precedent, or customer traction?",
            "Is the local-first architecture a competitive advantage (privacy) or disadvantage (no network effects, harder deployment) for scaling?",
            "What would you need to see in a competitive analysis to validate the claim that this market niche is unoccupied?",
        ],
        "For AI Governance Researchers": [
            "Does the ADDA metric capture what matters about automation bias, or does it oversimplify the phenomenon?",
            "The deviation documentation system forces the radiologist to explicitly reason about AI disagreement. Could this forced reflection actually change decision quality, creating a Hawthorne effect that confounds the research?",
            "What AI governance frameworks (e.g., NIST AI RMF, EU AI Act) would this documentation approach support or conflict with?",
        ],
        "For Quantitative Psychologists": [
            "What psychometric validation design would you recommend for the derived metrics — particularly ADDA, HRI, CPI, and the timing thresholds?",
            "The trust calibration probe (pre-AI expectation vs. actual reliance) is a single-item self-report measure. What established instrument should replace or supplement it?",
            "How many cases per reader are needed for stable estimates of individual ADDA rates? What's the minimum for meaningful between-reader comparison?",
        ],
        "For Human Factors / HCI Researchers": [
            "Is the keyboard shortcut system (0-5 for BI-RADS, Space to lock, +/- for confidence) an appropriate interaction design for high-stakes medical decisions, or does it trivialize the assessment?",
            "The workload thresholds (30/hr yellow, 40/hr red) are configurable but undocumented in their derivation. What human factors research should inform these values?",
            "Mouse hover as a proxy for visual attention — under what conditions is this valid, and what error bounds should be reported?",
        ],
        "For Biostatisticians": [
            "What MRMC statistical model (Obuchowski-Rockette, Hillis, etc.) is most appropriate for the proposed study design, given that both ADDA rate and timing metrics are outcomes of interest?",
            "The iMRMC export format (reads.csv + design.json) is implemented for FDA-standard analysis. What additional export formats would be needed for publication-quality analysis?",
            "With likely sample sizes of 10-30 radiologists and 20-50 cases, what effect sizes are detectable? Is this study powered to find clinically meaningful differences?",
        ],
        "For Malpractice Litigators (Plaintiff)": [
            "If opposing counsel presented a hash-chained decision log showing the radiologist spent 18 seconds before locking their impression, how would you attack it? What aspects of the documentation would you scrutinize most closely?",
            "Does the rubber-stamp detection system create a new liability — could a plaintiff argue that the system's own indicators flagged the radiologist as rushing?",
            "Would you hire an expert to challenge the hash chain's integrity, or would you focus your attack on the clinical judgment itself?",
        ],
        "For Malpractice Litigators (Defense)": [
            "How would you use this documentation in defending a radiologist who disagreed with the AI and was proven correct by pathology? How about one who agreed with the AI and was wrong?",
            "Is the temporal proof walkthrough (plain-language explanation of hash chain) effective for jury presentation, or is it too technical?",
            "What documentation would you need beyond what's in the expert witness packet to mount a complete defense?",
        ],
        "For Regulatory / Compliance Attorneys": [
            "What specific gaps exist between this platform's current state and HIPAA compliance? Is the 'no PHI leaves the device' architecture sufficient, or are there requirements beyond data locality?",
            "How does this documentation approach interact with state-specific medical malpractice reform statutes?",
            "If this were submitted as part of an FDA 510(k) for the documentation layer (not the AI itself), what class would it fall under? Does it qualify as a medical device?",
        ],
        "For Clinical Informatics Specialists": [
            "What FHIR resources would need to be mapped for EHR integration? Is DiagnosticReport sufficient, or would custom profiles be needed?",
            "How would the event stream integrate with existing clinical workflow engines (e.g., Epic Cognitive Computing, Cerner PowerChart)?",
            "What HL7 standards govern the communication of AI-assisted diagnostic findings, and does this system's event schema align with them?",
            "Would institutions be more likely to adopt this as a standalone tool or as an embedded module within their existing PACS/reporting system?",
        ],
    }

    for category, questions in panel_questions.items():
        doc.add_heading(category, level=3)
        for q in questions:
            p = doc.add_paragraph(style='List Number')
            p.add_run(q).font.size = Pt(10)

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Final line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of Document —")
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.size = Pt(10)
    run.italic = True

    return doc


if __name__ == '__main__':
    doc = create_document()

    primary = "./evidify-panel-review.docx"
    secondary = "./evidify-panel-review.docx"

    # Save primary
    os.makedirs(os.path.dirname(primary), exist_ok=True)
    doc.save(primary)
    print(f"Saved: {primary}")

    # Copy to secondary
    os.makedirs(os.path.dirname(secondary), exist_ok=True)
    shutil.copy2(primary, secondary)
    print(f"Copied: {secondary}")

    print("Done.")
